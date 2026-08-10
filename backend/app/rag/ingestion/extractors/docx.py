"""DOCX extraction.

Far easier than PDF: a .docx carries real structure. Paragraph styles name
headings outright, so none of the font-size guesswork in `pdf.py` is needed and
the section paths that come out are exact rather than inferred.

Word has no pages — pagination is computed by the renderer, not stored in the
file — so `page` stays `None` and citations fall back to the section path. Saying
that plainly is better than inventing a page number that would not match what the
user sees on screen.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

import docx
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.exceptions import DocumentProcessingError
from app.rag.ingestion.blocks import Block, BlockKind, ExtractedDocument

_HEADING_STYLE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)
_LIST_STYLE = re.compile(r"list|bullet", re.IGNORECASE)


def _heading_level(style_name: str) -> int | None:
    match = _HEADING_STYLE.match(style_name.strip())
    if match:
        return min(int(match.group(1)), 6)
    if style_name.strip().lower() == "title":
        return 1
    return None


def _iter_body(document: DocxDocument) -> list[Paragraph | Table]:
    """Paragraphs and tables in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate lists, which
    loses their relative order — a table would end up detached from the text that
    introduces it. Walking the body XML preserves it.
    """
    items: list[Paragraph | Table] = []
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            items.append(Paragraph(child, document))
        elif tag == "tbl":
            items.append(Table(child, document))
    return items


def _table_to_text(table: Table) -> str:
    """Render a table as pipe-delimited rows.

    Keeps the row/column association a model needs to read a figure out of a
    table, without pretending to reproduce the visual layout.
    """
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_docx(payload: bytes) -> ExtractedDocument:
    try:
        document: Any = docx.Document(BytesIO(payload))
    except Exception as exc:
        raise DocumentProcessingError("The Word document could not be read.") from exc

    blocks: list[Block] = []
    for item in _iter_body(document):
        if isinstance(item, Table):
            text = _table_to_text(item)
            if text:
                blocks.append(Block(text=text, kind=BlockKind.TABLE))
            continue

        text = item.text.strip()
        if not text:
            continue

        style_name = item.style.name if item.style is not None else ""
        level = _heading_level(style_name)
        if level is not None:
            blocks.append(Block(text=text, kind=BlockKind.HEADING, level=level))
        elif _LIST_STYLE.search(style_name):
            blocks.append(Block(text=text, kind=BlockKind.LIST_ITEM))
        else:
            blocks.append(Block(text=text, kind=BlockKind.PARAGRAPH))

    core_title = None
    try:
        core_title = (document.core_properties.title or "").strip() or None
    except Exception:  # pragma: no cover - malformed core properties
        core_title = None

    title = core_title or next((b.text for b in blocks if b.is_heading()), None)
    return ExtractedDocument(blocks=blocks, page_count=None, title=title)
